from pathlib import Path
from docxtpl import DocxTemplate
import openpyxl
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.shared import Cm
import os
import numpy as np
import re
import datetime
# from PyInquirer import prompt
import pprint
import Helfer_Objekte
from  Helfer_Objekte import check_invoice_archive,question_next_invoice_number
import dateutil.parser
import tkinter as tk

def make_invoice_praxis(allhourdata_path,allclientdata_path,supparentdir,excel_template_path,template_path):
    # read in
    allhourdata = pd.read_excel(allhourdata_path, parse_dates=[0])
    allclientdata = pd.read_excel(allclientdata_path, index_col=0, header=None, sheet_name=None)


    #select which client
    allclientsnames = list(allclientdata.keys())
    allclientsnames.sort()
    newclienttext = "Keiner der obigen Personen -> mach neue Person"
    allclientsnames.append(newclienttext)


    clientname = Helfer_Objekte.ask_many_multiple_choice_question(
        "Von welcher Person willst du die Rechnung ausdrucken?",
        allclientsnames
    )
    print("Ich mache die Rechnung für die Person:")
    print(clientname)
    newclient = False
    if clientname == newclienttext:
        newclient = True
        #execute the input Routine
        print("Dann lass uns einen Eintrag in der KlientInnendaten Datei anlegen: ")
        clientdata = Helfer_Objekte.input_new_person(allclientdata_path)
        clientname = clientdata["name"]
        print("Dann können wir auch hier gleich die letzten Stundendaten eintrage:")
        namehourdata = Helfer_Objekte.insert_hourdata(allhourdata_path, clientname)
        namehourdata["Minuten"] = [int(x) for x in namehourdata["Minuten"]]



    # prepare data
    if newclient == False:
        namehourdata = allhourdata[allhourdata["Name"] == clientname] #select only the hours of the given name
        #namehourdata["Datum"] = list(map(lambda x: dateutil.parser.parse(x), namehourdata["Datum"]))


    try:
        outputdir_path = os.path.join(supparentdir, f"{datetime.datetime.today().year}")
        archive_which_invoices_path = os.path.join(outputdir_path, f"Rechnungen {datetime.datetime.today().year}.xlsx")
        archive =pd.read_excel(archive_which_invoices_path)
        invoicetime_last = archive.iloc[:,1][archive.iloc[:,2] == clientname].values[-1]
        print("got last invoice times")
    except:
        print("no last invoice times")
        invoicetime_last = 0
    #select with a calendar
    invoice_start_date, invoice_end_date = Helfer_Objekte.get_date(namehourdata,invoicetime_last)
    invoice_start_date = pd.to_datetime(invoice_start_date)
    invoice_end_date = pd.to_datetime(invoice_end_date)
    # invoice_start_date, invoice_end_date = datetime.datetime(2023,1,6), datetime.datetime(2023,2,27)

    print("Ich nehme alle Termine von " + clientname + " ab: " + invoice_start_date.strftime("%d.%m.%Y") + " bis zum " + invoice_end_date.strftime("%d.%m.%Y") )

    namehourdata = namehourdata[(namehourdata.Datum >= invoice_start_date)&(namehourdata.Datum <= invoice_end_date)]   #delete everything brfore lastinvoicegroup

    # get the invoice archive or make new archive
    year_of_invoice = invoice_end_date.year
    invoicenumber_pattern = r'(\d{4})-(\d+)'
    lastinvoice_num = check_invoice_archive(year_of_invoice,supparentdir,excel_template_path,invoicenumber_pattern= invoicenumber_pattern)

    thisinvoicenumber = question_next_invoice_number(year_of_invoice,lastinvoice_num,invoicenumber_pattern)

    if not os.path.isdir(outputdir_path):
        os.mkdir(outputdir_path)
    else:
        print(f"We already have a directory {outputdir_path}")


    outputfile_path = os.path.join(outputdir_path, f"RE {thisinvoicenumber} {clientname} {datetime.date.today().strftime('%d_%m_%Y')}.docx")

    print(f"Now i can create the outputdata filepaths:   \n{archive_which_invoices_path}\n{outputfile_path}")

    if newclient == False:

        clientdata = allclientdata[clientname].to_dict()[1]
        print("Die Patientendaten sind:")
        pprint.pprint(clientdata)

    #additional data on invoice
    if clientdata["Kind"] == "nein":
        clientdata["BeideElternteile"] = clientdata["Name"]
        clientdata
    elif clientdata["Kind"] == "ja":
        clientdata["BeideElternteile"] = str(clientdata["Elternteil1"]) + " und " + str(clientdata["Elternteil2"])
    else:
        print("Nicht gegeben ob Kind oder Erwachsen")
        raise SystemExit

    clientdata["Rechnungsnummer"] = thisinvoicenumber
    clientdata["Heute"] = datetime.date.today().strftime("%d.%m.%Y")
    if clientdata["Kind"] == "ja":
        if clientdata["Geschlecht"] == "m":
            clientfirstname = clientdata["Name"].split()[0]
            clientdata["Wordkindtext"] = " für Ihren Sohn " + clientfirstname + ", geboren am " + clientdata["Geb."].strftime("%d.%m.%Y") + ","
        if clientdata["Geschlecht"] == "w":
            clientfirstname = clientdata["Name"].split()[0]
            clientdata["Wordkindtext"] = " für Ihre Tochter " + clientfirstname + ", geboren am " + clientdata["Geb."].strftime("%d.%m.%Y") + ","



    # input the client data  in word
    doc = DocxTemplate(template_path)
    doc.render(clientdata)
    doc.save(outputfile_path)

    ## input the hour table in word
    doc = Document(outputfile_path)
    doc.tables #a list of all tables in document
    # table nr. 0 is the data table and table nr. 1 is the sum table

    # prepare the hour table
    amountpersession = namehourdata["Minuten"].apply(lambda x: round(x * float(clientdata["Stundensatz"])/60, 1))
    amountpersession = amountpersession.rename("Betrag_pro_Einheit")

    wordtable = pd.concat([namehourdata["Datum"].apply(lambda x: x.strftime("%d.%m.%Y")), namehourdata["Minuten"].apply(lambda x: str(x) + " min"), amountpersession.apply(lambda x: "%0.2f" % x + " €") ], axis=1)
    print("Die Stunden sind: " )
    print(wordtable)


    # insert the table in the Word document
    for index, row in wordtable.iterrows():
        hourdatatable = doc.tables[0]   #so hourdatatable is the first table in the document
        data_row = hourdatatable.add_row().cells
        for i,(name,entry) in enumerate(row.items()):
                data_row[i].text = entry
    #format it
    for row in doc.tables[0].rows:
        row.height = Cm(0.8)
        row.alignment = WD_TABLE_ALIGNMENT.CENTER


    #insert total amount into tables[1]
    totalamount = sum(np.array(amountpersession))
    totalamountstring = (str(totalamount)+"0").replace(".",",")
    doc.tables[1].cell(0, 2).text = str(totalamount) + "0" + " €"

    clientdata["Stundeninfo"] = wordtable
    save_or_not = Helfer_Objekte.ask_to_save_with_dict(clientdata)
    print(save_or_not)
    if save_or_not:
        print(f"Save word file to {outputfile_path}")
        doc.save(outputfile_path)

        #write what I did in the archive
        try_saving = True
        while try_saving:
            try:
                Helfer_Objekte.save_to_archive(thisinvoicenumber,clientdata["Heute"],clientname,invoice_start_date,invoice_end_date,totalamount,archive_which_invoices_path)
                try_saving = False
            except:
                def show_alert():
                    # Display an alert message box with an "Okay" button
                    tk.messagebox.showinfo("Fehler", f"Ich konnte die Excel nicht speichern. Schließe zuerst die Datei {archive_which_invoices_path}")
                print("error in saving archive")
                root = tk.Tk()
                root.withdraw()  # Hide the main window

                # Show the alert
                show_alert()

                # Continue with the rest of the code after the alert
                print("The alert has been dismissed. Continuing with the rest of the code...")

                # Close the Tkinter application
                root.quit()

        print(f"\n Die Rechnung wurde in einem Word Dokument erstellt, zu finden unter Desktop -> Rechnungen-Verknüpfung -> Jahr {year_of_invoice}")


        os.startfile(archive_which_invoices_path)
        os.startfile(outputfile_path)
    else:
        print("Didnot save anything")