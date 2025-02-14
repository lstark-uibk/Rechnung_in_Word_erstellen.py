from docxtpl import DocxTemplate
import openpyxl
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.shared import Cm
import os
import numpy as np
import datetime
import pprint
from Helfer_Objekte import (check_invoice_archive,question_next_invoice_number, select_client, get_date, save_to_archive, ask_to_save,
                                     stringsandyear_topath, stringsandinvoicenumber_topath)
import tkinter as tk


def make_invoice_praxis(allhourdata_path,allclientdata_path,excel_template_path,template_path,outputdir_suppath,
                        nameouptutdir,nameoutputarchivefile, invoicenumber_pattern, invoicenumber_pattern_names, nameinvoicefile,user = "r"):
    ### parameters:



    #readin
    allhourdata = pd.read_excel(allhourdata_path, parse_dates=[0])
    allclientdata = pd.read_excel(allclientdata_path, index_col=0, header=None, sheet_name=None)


    #select which client
    allclientsnames = list(allclientdata.keys())
    allclientsnames.sort()
    clientname = select_client(
        allclientsnames
    )
    print("Ich mache die Rechnung für die Person:")
    print(clientname)
    namehourdata = allhourdata[allhourdata["Name"] == clientname]

    # check whether there was already an invoice for this person in the last 3 years (we want to see it in the calendar)
    invoicetime_last = 0
    for yearback in [2,1,0]:
        yearcheck = datetime.datetime.today().year - yearback
        directory_to_search = stringsandyear_topath(nameouptutdir,yearcheck)
        fullpath_to_search = os.path.join(outputdir_suppath,directory_to_search)
        try:
            archive_which_invoices_path = os.path.join(fullpath_to_search, f"Rechnungen {yearcheck}.xlsx")
            print(f"Check {archive_which_invoices_path}")
            archive =pd.read_excel(archive_which_invoices_path)
            invoicetime_last = archive.iloc[:,1][archive.iloc[:,2] == clientname].values[-1]
            print(f"got last invoice time in year {yearcheck}: {invoicetime_last}")
        except:
            print(f"no last invoice times in year {yearcheck}")

    #select with a calendar
    invoice_start_date, invoice_end_date = get_date(namehourdata,invoicetime_last)
    invoice_start_date = pd.to_datetime(invoice_start_date)
    invoice_end_date = pd.to_datetime(invoice_end_date)
    print("Ich nehme alle Termine von " + clientname + " ab: " + invoice_start_date.strftime("%d.%m.%Y") + " bis zum " + invoice_end_date.strftime("%d.%m.%Y") )
    namehourdata = namehourdata[(namehourdata.Datum >= invoice_start_date)&(namehourdata.Datum <= invoice_end_date)]   #delete everything brfore lastinvoicegroup

    # now fix the year of which the invoice is made
    year_of_invoice = invoice_end_date.year

    #now fix the output path and check whether it exists
    outputdir = stringsandyear_topath(nameouptutdir, year_of_invoice)
    outputdir_path = os.path.join(outputdir_suppath, outputdir)
    if not os.path.isdir(outputdir_path):
        os.mkdir(outputdir_path)
    else:
        print(f"We already have a output directory {outputdir_path}")

    #now get the invoice archive or make new archive and check for invoice numbers
    archive_which_invoices_name = stringsandyear_topath(nameoutputarchivefile,year_of_invoice)
    archive_which_invoices_path = os.path.join(outputdir_path, archive_which_invoices_name)
    print(f"Year to link this invoice to: {year_of_invoice}")
    lastinvoice_num = check_invoice_archive(year_of_invoice, outputdir_path, archive_which_invoices_path,
                                            excel_template_path, invoicenumber_pattern= invoicenumber_pattern)
    print(f"lastinvoice_num: {lastinvoice_num}")

    # check whether invoice number is okay
    thisinvoicenumber = question_next_invoice_number(year_of_invoice,lastinvoice_num,invoicenumber_pattern,invoicenumber_pattern_names)

    # since we now have the year and the invoicenumber we set outputfilepath
    filename = stringsandinvoicenumber_topath(nameinvoicefile,thisinvoicenumber,clientname, datetime.date.today().strftime('%d_%m_%Y'))
    outputfile_path = os.path.join(outputdir_path, filename)
    print(f"Now i can create the outputdata filepaths:   \n{archive_which_invoices_path}\n{outputfile_path}")

    # data processing
    clientdata = allclientdata[clientname].to_dict()[1]
    print("Die Patientendaten sind:")
    pprint.pprint(clientdata)

    #additional data on invoice
    if clientdata["Kind"] == "nein":
        clientdata["BeideElternteile"] = clientdata["Name"]
        clientdata
    elif clientdata["Kind"] == "ja":
        clientdata["BeideElternteile"] = str(clientdata["Elternteil1"]) + " und " + str(clientdata["Elternteil2"])
    else:
        print("Nicht gegeben ob Kind oder Erwachsen")
        raise SystemExit

    clientdata["Rechnungsnummer"] = thisinvoicenumber
    clientdata["Heute"] = datetime.date.today().strftime("%d.%m.%Y")
    if clientdata["Kind"] == "ja":
        if clientdata["Geschlecht"] == "m":
            clientfirstname = clientdata["Name"].split()[0]
            clientdata["Wordkindtext"] = " für Ihren Sohn " + clientfirstname + ", geboren am " + clientdata["Geb."].strftime("%d.%m.%Y") + ","
        if clientdata["Geschlecht"] == "w":
            clientfirstname = clientdata["Name"].split()[0]
            clientdata["Wordkindtext"] = " für Ihre Tochter " + clientfirstname + ", geboren am " + clientdata["Geb."].strftime("%d.%m.%Y") + ","


    #preprocessdata
    amountpersession = namehourdata["Minuten"].apply(lambda x: round(x * float(clientdata["Stundensatz"]) / 60, 1))
    amountpersession = amountpersession.rename("Betrag_pro_Einheit")

    wordtable = pd.concat([namehourdata["Datum"].apply(lambda x: x.strftime("%d.%m.%Y")),
                           namehourdata["Minuten"].apply(lambda x: str(x) + " min"),
                           amountpersession.apply(lambda x: "%0.2f" % x + " €")], axis=1)
    print("Die Stunden sind: ")
    print(wordtable)

    # insert total amount into tables[1]
    totalamount = sum(np.array(amountpersession))
    clientdata["Stundeninfo"] = wordtable
    clientdata_list = [[key,clientdata[key]] for key in clientdata]


    # process to prepare output for user r and b
    if user == "r":
        # input the client data  in word
        doc = DocxTemplate(template_path)
        doc.render(clientdata)
        doc.save(outputfile_path)

        ## input the hour table in word
        doc = Document(outputfile_path)
        doc.tables


        # insert the table in the Word document
        for index, row in wordtable.iterrows():
            hourdatatable = doc.tables[0]   #so hourdatatable is the first table in the document
            data_row = hourdatatable.add_row().cells
            for i,(name,entry) in enumerate(row.items()):
                    data_row[i].text = entry
        #format it
        for row in doc.tables[0].rows:
            row.height = Cm(0.8)
            row.alignment = WD_TABLE_ALIGNMENT.CENTER



        totalamountstring = (str(totalamount)+"0").replace(".",",")
        doc.tables[1].cell(0, 2).text = str(totalamount) + "0" + " €"


    if user == "b":
        invoice = openpyxl.load_workbook(template_path)
        invoice_sheet = invoice ['Rechnung']
        excelsheet_locs = {"Name": ("B", 10),
                           "Adresse": ("B", 12),
                           "Stadt": ("B", 13),
                           "Heute": ("J", 14),
                           "Rechnungsnummer": ("J", 15),
                           "Versicherungsnummer": ("J", 17)}
        usevalues = ["Name", "Adresse", "Stadt", "Heute", "Rechnungsnummer", "Versicherungsnummer"]
        for value in usevalues:
            location = f"{excelsheet_locs[value][0]}{excelsheet_locs[value][1]}"
            invoice_sheet[location] = f"{clientdata[value]}"

        firstrows_hourdata = {"Datum": ("C", 22), "Anzahl": ("E", 22), "Preis/Einh.": ("G", 22)}

        i = 0
        for row, session in namehourdata.iterrows():
            invoice_sheet[f"{firstrows_hourdata['Datum'][0]}{firstrows_hourdata['Datum'][1] + i}"] = session["Datum"]
            invoice_sheet[f"{firstrows_hourdata['Anzahl'][0]}{firstrows_hourdata['Anzahl'][1] + i}"] = session[
                                                                                                           "Minuten"] / 60
            invoice_sheet[f"{firstrows_hourdata['Preis/Einh.'][0]}{firstrows_hourdata['Preis/Einh.'][1] + i}"] = \
            clientdata["Stundensatz"]
            i += 1

    # now ask to save
    save_or_not = ask_to_save(clientdata_list)
    print(f"save or not {save_or_not}")
    if save_or_not:

        if user =="r":
            print(f"Save word file to {outputfile_path}")
            doc.save(outputfile_path)
        if user == "b":
            invoice.save(outputfile_path)
            print(f"Save excel file to {outputfile_path}")

        #write what I did in the archive
        try_saving = True
        while try_saving:
            try:
                save_to_archive(thisinvoicenumber,clientdata["Heute"],clientname,invoice_start_date,invoice_end_date,totalamount,archive_which_invoices_path)
                try_saving = False
            except Exception as e:
                def show_alert():
                    # Display an alert message box with an "Okay" button
                    tk.messagebox.showinfo("Fehler", f"Ich konnte die Excel nicht speichern. Schließe zuerst die Datei {archive_which_invoices_path}")
                print("error in saving archive")
                print(e)
                root = tk.Tk()
                root.withdraw()  # Hide the main window

                # Show the alert
                show_alert()

                # Continue with the rest of the code after the alert
                print("The alert has been dismissed. Continuing with the rest of the code...")

                # Close the Tkinter application
                root.quit()

        print(f"\n Die Rechnung wurde in einem Word Dokument erstellt, zu finden unter Desktop -> Rechnungen-Verknüpfung -> Jahr {year_of_invoice}")

        if os.name == 'posix':
            print("This system is Linux or another Unix-like system.")
            import subprocess
            subprocess.run(['xdg-open', archive_which_invoices_path])
            subprocess.run(['xdg-open', outputfile_path])

        else:
            print("This system is not Linux.")
            os.startfile(archive_which_invoices_path)
            os.startfile(outputfile_path)
    else:
        print("Didnot save anything")
    from Rechnung_erstellen import main
    main()